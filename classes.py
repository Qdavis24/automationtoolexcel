import pandas
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import yaml
import re


def extract_section_number(text):
    """Extract section number from text using regex pattern.

        Args:
            text: Input text to extract section number from.

        Returns:
            Extracted section number in X.XX format.

        Raises:
            Exits script if no section number found.
        """
    pattern = r'\d\.\d{2}'
    match = re.search(pattern, text)
    number = match.group() if match else None
    if not number:
        input(f"Could not extract section number from: {text}...\n"
              f"Possible formatting changes to template may have broken script...\n"
              f"Please retrieve approved word template for this automation tool")
        exit(1)
    return number


class ConfigError(Exception):
    pass


class WorksheetError(Exception):
    pass


class Config:
    def __init__(self, filepath):
        """Initialize Config object by loading and validating YAML configuration.

           Args:
               filepath: Path to the YAML configuration file.

           Attributes loaded:
               - Excel filepaths (data, template, export)
               - Spreadsheet column configurations
               - Advanced Excel parsing settings

           Raises:
               Exits script on YAML parsing or file access errors.
           """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
            if 'excel' not in self.config or 'filepaths' not in self.config:
                raise ConfigError("Missing required sections in config file")
        except yaml.YAMLError as error:
            input(f"Invalid YAML format in config file: {error}")
            exit(1)
        except FileNotFoundError as error:
            print("\n----------------ERROR--------------------\n")
            input(f"{error}: config file 'config.yaml' not found.")
            exit(1)
        else:
            # load filepaths
            self.excel_filepath = self._verify_extensions(self.config['filepaths']['data_spreadsheet'], ".xlsx")
            self.word_template_filepath = self._verify_extensions(self.config['filepaths']['document_template'],
                                                                  ".docx")
            self.word_export_filepath = self._verify_extensions(self.config['filepaths']['final_document'], ".docx")

            # load typical config data
            self.standard_col = self.config['excel']['standard_column']
            self.rfi_col = self.config['excel']['rfi_column']
            self.sheet_name = self.config['excel']['sheet_name']

            # load advanced config data
            self.row_shift = self.config['excel']['advanced']['row_shift']
            self.header = self.config['excel']['advanced']['header']
            self.ignore_color = self.config['excel']['advanced']['ignore_color']

    def _verify_extensions(self, filepath: str, ext: str):
        """Validate and correct file extension.

           Args:
               filepath: Path to the file.
               ext: Expected file extension.

           Returns:
               Filepath with corrected extension.
           """
        split = filepath.lstrip(".").split(".")
        if len(split) < 2:
            return "." + split[0] + ext
        if split[1] != ext:
            return "." + split[0] + ext

    def _index_from_letter(self, char):
        """Convert Excel column letter to zero-based index.

           Args:
               char: Single column letter (A-Z).

           Returns:
               Zero-based column index.

           Raises:
               ConfigError: If input is not a valid A-Z column letter.
           """
        ascii_value = ord(char.upper())
        if ascii_value < 65 or ascii_value > 90:
            raise ConfigError(f"Invalid column letter")
        return ascii_value - 65

    @property
    def rfi_col_index(self):
        """Convert RFI column letter to zero-based index.

           Returns:
               Zero-based column index for RFI column.

           Raises:
               SystemExit: If RFI column letter is invalid.
           """
        try:
            return self._index_from_letter(self.rfi_col)
        except ConfigError as error:
            input(f"{error}: Invalid charactor used for RFI column letter - ({self.rfi_col})...\n"
                  f"please use a-z or A-Z")
            exit(1)

    @property
    def standard_col_index(self):
        """Convert standard column letter to zero-based index.

           Returns:
               Zero-based column index for standard column.

           Raises:
               SystemExit: If standard column letter is invalid.
           """
        try:
            return self._index_from_letter(self.standard_col)
        except ConfigError as error:
            input(f"{error}: Invalid charactor used for RFI column letter - ({self.standard_col})...\n"
                  f"please use a-z or A-Z")
            exit(1)


class ExcelDf:
    KEYWORD_STRING = "see"
    KEYWORD_SLICE = (0, 3)
    RFI_COLUMN_NAME = "RFI"
    STANDARD_COLUMN_NAME = "STANDARD"

    def __init__(self, filepath: str, sheet_name: str, row_shift: int, header: int):
        """Initialize ExcelDf object by loading Excel spreadsheet.

           Args:
               filepath: Path to the Excel file.
               sheet_name: Name of the worksheet to load.
               row_shift: Number of rows to shift during processing.
               header: Row number containing column headers.

           Attributes:
               df: Pandas DataFrame containing Excel sheet data.
               worksheet: Active worksheet from loaded workbook.
               ROW_SHIFT: Configured row shift value.

           Raises:
               SystemExit: On worksheet access, sheet name, or file path errors.
           """
        try:
            self.df = pd.read_excel(filepath, sheet_name=sheet_name, header=header)
            self.worksheet = load_workbook(filepath).active
            if self.worksheet is None:
                raise WorksheetError("No active worksheet found")
        except WorksheetError as error:
            input(f"{error}: Error accessing worksheet...\n"
                  f"Excel file may be corrupted")
            exit(1)
        except ValueError as error:
            input(f"{error}: Excel sheet name specified is incorrect.")
            exit(1)

        except FileNotFoundError as error:
            input(f"{error}: Excel file path specified is incorrect.")
            exit(1)
        else:
            self.ROW_SHIFT = row_shift
            print("\n----------------EXCEL DATA PREVIEW--------------------\n")
            print(f"Filepath to xlsx: {filepath}")
            print(f"Name of sheet: {sheet_name}")
            print(f"Number of rows: {self.df.shape[0]}\n"
                  f"Number of columns: {self.df.shape[1]}\n")
            print(f"Column names {self.df.columns.tolist()}")

            print("\n------------------------------------------------------\n")

    def check_cell_fill(self, col: str):
        """ prints out the argb values for fill in each cell in a specified col """
        for cell in self.worksheet[col]:
            print(cell.fill.start_color.rgb)

    def compress_df(self, standard_col_index: int, rfi_col_index: int):
        """Reduce DataFrame to specified columns and rename them.

            Args:
                standard_col_index: Column index for standard column.
                rfi_col_index: Column index for RFI column.

            Raises:
                SystemExit: If specified column indexes are out of range.
            """
        try:
            self.df = self.df.iloc[:, [standard_col_index, rfi_col_index]]
            self.df.columns = [self.STANDARD_COLUMN_NAME, self.RFI_COLUMN_NAME]
        except IndexError as error:
            print("\n----------------ERROR--------------------\n")
            input(f"{error}: Indexes specified for columns are incorrect\n"
                  f"EXISTING COLUMNS..\n"
                  f"{self.df.columns}")
            exit(1)
        else:
            print("Successfully compressed dataframe...")

    def clean_df(self, color_value, color_col_char):
        self._remove_rows(self._gen_drop_indexes_color(color_value, color_col_char))
        self._remove_na()
        self._remove_see_from_rfi()

    def _remove_na(self):
        self.df = self.df.dropna()
        print("Successfully dropped NA rows from dataframe...")

    def _remove_see_from_rfi(self):
        """Remove rows from DataFrame where RFI column starts with 'see'.

           Filters out rows in the RFI column that begin with 'see' (case-insensitive),
           effectively cleaning the DataFrame by excluding entries marked for reference.
           """
        self.df = self.df[~self.df[self.RFI_COLUMN_NAME].astype(str).str.lower().str.startswith('see')]

    def _gen_drop_indexes_color(self, argb: str, col_char) -> list:
        """Generate list of row indexes with specific background color.

           Searches specified Excel column for cells with matching background color.

           Args:
               argb: ARGB color value to match.
               col_char: Excel column letter to search.

           Returns:
               List of row indexes with matching background color,
               adjusted to match DataFrame row indexes.
           """
        return [cell.row - self.ROW_SHIFT for cell in self.worksheet[col_char] if cell.fill.start_color.rgb == argb]

    def _remove_rows(self, remove_indexes: list):
        """Remove specified rows from DataFrame.

           Args:
               remove_indexes: List of row indexes to drop.

           Raises:
               SystemExit: If row indexes are invalid or out of range.
           """
        try:
            self.df = self.df.drop(remove_indexes)
        except IndexError as error:
            print("\n----------------ERROR--------------------\n")
            input(f"{error}: Row shift index configuration is incorrect\n")
            exit(1)
        else:
            print("Successfully removed colored rows from dataframe...")

    def display_dataframe(self):
        print("\n----------------EXCEL DATA CLEANED--------------------\n")
        print(self.df.head())
        print("\n------------------------------------------------------\n")

    def uniques(self, column: str) -> list:
        return self.df[column].unique().tolist()

    def return_series(self, column: str) -> pandas.Series:
        return self.df[column]


class Data:
    def __init__(self):
        self.mapped_questions = {}

    def map_questions_to_sections(self, sections: list, df: pandas.DataFrame, standard_col_name: str,
                                  rfi_col_name: str):
        """Create dictionary mapping section numbers to their corresponding questions.

           Args:
               sections: Unique values representing section titles.
               df: DataFrame containing section and question data.
               standard_col_name: Column name for section identifiers.
               rfi_col_name: Column name containing questions.
           """
        for section in sections:
            section_qs = [str(question) for question in df[df[standard_col_name] == section][rfi_col_name].tolist()]
            if section_qs:
                self.mapped_questions[extract_section_number(section)] = section_qs

    def pretty_print(self):
        print("\n----------------MAPPED DATA PREVIEW--------------------\n")
        for key, value in self.mapped_questions.items():
            print(f"{key}: ")
            num = 0
            if type(value) != list:
                print(value)
            else:
                for q in value:
                    num += 1
                    print(f"{num}) {q.strip()}")
            print("\n")
        print("\n------------------------------------------------------\n")


class Word:
    START_SLICE = (14, 15)
    START_STRING = "–"
    INSERT_SLICE = (0, -1)
    INSERT_STRING = "Confirm/Submit/Describe"

    def __init__(self, filepath: str, mapped_questions: dict):
        """Initialize Word document processing object.

           Args:
               filepath: Path to the Word document template.
               mapped_questions: Dictionary mapping section numbers to questions.

           Attributes:
               doc: Loaded Word document.
               mapped_questions: Input dictionary of section questions.
               to_delete: List of paragraph indexes to remove.
               to_modify: List of paragraph indexes to modify.
           """
        try:
            self.doc = Document(filepath)
        except PackageNotFoundError as error:
            print("\n----------------ERROR--------------------\n")
            input(f"{error}: word template file path specified is incorrect.")
            exit(1)
        else:
            self.mapped_questions = mapped_questions
            self.to_delete = []
            self.to_modify = []

    def _gen_modify_indexes(self):
        """ generates all par indexes for our to be modified document paragraphs """
        curr_section = None
        for i, par in enumerate(self.doc.paragraphs):
            if self._check_for_question(par):

                if curr_section:
                    self.to_modify.append((curr_section, i))

            if self._check_for_section_head(par):
                identifier = extract_section_number(par.text)
                curr_section = identifier if identifier in self.mapped_questions else False

    def _gen_delete_indexes(self):
        """ generates all par indexes for our to be deleted document paragraphs """
        delete = False
        for i, par in enumerate(self.doc.paragraphs):
            if self._check_for_section_head(par):
                if delete:
                    self.to_delete.append((start, i - 1))
                start = i
                identifier = extract_section_number(par.text)
                delete = identifier not in self.mapped_questions

    def _gen_paragraph(self, par, text):
        par.text = f"{text}\n\n[enter response here]\n\n"
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = 0

        return par._element

    def modify(self):
        """ goes through the modify indexes and then inserts the new content using the map """
        self._gen_modify_indexes()
        for par_index_section_key in self.to_modify[::-1]:
            curr_p = self.doc.paragraphs[par_index_section_key[1]]
            curr_p.clear()
            curr_p = self._gen_paragraph(curr_p, self.mapped_questions[par_index_section_key[0]][0])

            for question in self.mapped_questions[par_index_section_key[0]][1:]:
                new_p = self.doc.add_paragraph()
                curr_p.addnext(self._gen_paragraph(new_p, question))

    def remove(self):
        """ goes through delete indexes and then deletes those paragraphs from document """
        self._gen_delete_indexes()
        for tup in self.to_delete[::-1]:
            for i in range(tup[1], tup[0] - 1, -1):
                p = self.doc.paragraphs[i]._element
                p.getparent().remove(p)

    def _check_for_section_head(self, par):
        """ checks for par that represents a new section in word doc
        ARGS:
        -------------------------
            par: the current par we are checking
        RETURN:
        -------------------------
            True if it is a section head
            False if not a section head
            """

        if par.text[slice(*self.START_SLICE)] == self.START_STRING:
            return True
        return False

    def _check_for_question(self, par):
        """ checks for par that represent where we should have our new questions inserted at in word doc
        ARGS:
        -------------------------
            par: the current par we are checking
        RETURN:
        -------------------------
            True if it is an insert questions par
            False if not an insert questions par
            """
        if par.text[slice(*self.INSERT_SLICE)] == self.INSERT_STRING:
            return True
        return False

    def check_modify_indexes(self):
        """ prints out all sections with their current document text """
        for tup in self.to_modify:
            print(tup[0], self.doc.paragraphs[tup[1]].text)
            print("\n-------------------next-------------------")

    def check_delete_indexes(self):
        """ prints out all document paragraph text within delete indexes """
        for tup in self.to_delete:
            for i in range(tup[0], tup[1] + 1):
                print(self.doc.paragraphs[i].text)
            print("\n-------------------next-------------------")

    def save(self, filepath: str):
        try:
            self.doc.save(filepath)
        except Exception as error:
            input(f"{error}: Error saving document")
            exit(1)
